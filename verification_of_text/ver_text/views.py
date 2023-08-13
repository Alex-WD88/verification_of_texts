from django.shortcuts import render, redirect
from .models import File, ChildFile
import difflib
import chardet
from PyPDF2 import PdfFileReader
import textract
from docx import Document
from openpyxl import load_workbook
from . import models


def upload(request):
    if request.method == 'POST':
        new_file = models.File()
        new_file.document = request.FILES['document']
        new_file.save()
        return redirect('/')
    return render(request, 'ver_text/index.html')


def file(request, id):
    file = models.File.objects.get(id=id)
    child_files = models.ChildFile.objects.filter(parent=file)

    if request.method == 'POST':
        child_files = models.ChildFile()
        child_files.parent = file

        if 'document' in request.FILES:
            child_files.document = request.FILES['document']

        child_files.save()
        return redirect('file', id=file.id)

    return render(request, 'ver_text/file.html', {'file': file, 'child_files': child_files})


def all_file(request):
    files = models.File.objects.all()
    return render(request, 'ver_text/allFiles.html', {'files': files})


def all_child_files(request):
    child_files = models.ChildFile.objects.all()
    return render(request, 'ver_text/allchildfiles.html', {'child_files': child_files})


def compare(request, id):
    file = File.objects.get(id=id)
    child_files = ChildFile.objects.filter(parent=file)

    diffs = []
    for child_file in child_files:
        diff = compare_files(file.document.path, child_file.document.path)
        diffs.append(diff)

    diff = {
        0: {'equal': 'Совпадение'},
        1: {'replace': 'Замена'}
    }

    return render(request, 'ver_text/compare.html',
                  {'file': file, 'child_files': child_files, 'diffs': diffs, 'diff': diff})


def compare_files(file1_path, file2_path):
    file1_extension = file1_path.split('.')[-1].lower()
    file2_extension = file2_path.split('.')[-1].lower()

    if file1_extension == 'txt' or file2_extension == 'txt':
        # Если один из файлов - текстовый файл, использовать существующую логику
        return compare_text_files(file1_path, file2_path)
    elif file1_extension == 'pdf' and file2_extension == 'pdf':
        # Если оба файла являются PDF-файлами, сравнить их содержимое
        return compare_pdf_files(file1_path, file2_path)
    elif file1_extension in ['doc', 'docx'] and file2_extension in ['doc', 'docx']:
        # Если оба файла являются Word-документами, сравнить их содержимое
        return compare_word_documents(file1_path, file2_path)
    elif file1_extension in ['xls', 'xlsx'] and file2_extension in ['xls', 'xlsx']:
        # Если оба файла являются Excel-файлами, сравнить их содержимое
        return compare_excel_files(file1_path, file2_path)
    else:
        return "Unsupported file format"


def compare_text_files(file1_path, file2_path):
    # Определить кодировку файлов
    file1_encoding = detect_encoding(file1_path)
    file2_encoding = detect_encoding(file2_path)

    # Открыть и прочитать оба файла с определенной кодировкой
    with open(file1_path, 'r', encoding=file1_encoding) as file1, open(file2_path, 'r', encoding=file2_encoding) as file2:
        file1_content = file1.readlines()
        file2_content = file2.readlines()

    # Использовать difflib для сравнения содержимого строк-по-строчно
    diff = difflib.Differ()
    diff_result = list(diff.compare(file1_content, file2_content))

    return diff_result


def compare_pdf_files(file1_path, file2_path):
    # Извлечь текст из PDF-файлов
    text1 = textract.process(file1_path).decode('utf-8')
    text2 = textract.process(file2_path).decode('utf-8')

    # Использовать difflib для сравнения содержимого PDF-файлов
    diff = difflib.Differ()
    diff_result = list(diff.compare(text1.splitlines(), text2.splitlines()))

    return diff_result


def compare_word_documents(file1_path, file2_path):
    # Открыть оба Word-документа
    doc1 = Document(file1_path)
    doc2 = Document(file2_path)

    # Получить содержимое параграфов Word-документов
    text1 = ''
    text2 = ''
    for para in doc1.paragraphs:
        text1 += para.text
    for para in doc2.paragraphs:
        text2 += para.text

    # Использовать difflib для сравнения содержимого Word-документов
    diff = difflib.Differ()
    diff_result = list(diff.compare(text1.splitlines(), text2.splitlines()))

    return diff_result


def compare_excel_files(file1_path, file2_path):
    # Открыть оба Excel-файла
    workbook1 = load_workbook(file1_path)
    workbook2 = load_workbook(file2_path)

    # Получить содержимое строк Excel-файлов
    rows1 = []
    rows2 = []
    for sheet in workbook1.sheetnames:
        for row in workbook1[sheet].iter_rows(values_only=True):
            rows1.append(row)
    for sheet in workbook2.sheetnames:
        for row in workbook2[sheet].iter_rows(values_only=True):
            rows2.append(row)

    # Использовать difflib для сравнения содержимого Excel-файлов
    diff = difflib.Differ()
    diff_result = list(diff.compare([str(row) for row in rows1], [str(row) for row in rows2]))

    return diff_result


def detect_encoding(file_path):
    with open(file_path, 'rb') as file:
        byte_data = file.read()

    result = chardet.detect(byte_data)
    encoding = result['encoding']
    confidence = result['confidence']

    if not encoding or confidence < 0.8:
        encoding = 'utf-8'  # Используйте кодировку utf-8 по умолчанию

    return encoding
